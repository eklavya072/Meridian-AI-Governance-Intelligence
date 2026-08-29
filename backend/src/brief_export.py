"""Executive Brief exports — DOCX (python-docx) and PDF (reportlab).

Both render from the SAME stored structured brief dict (reports.meta) so the
two formats never drift. Each includes a simple branded header/footer:
Meridian mark, generation date, page number.

Why reportlab for PDF instead of weasyprint: weasyprint needs pango/cairo
system libraries that are not guaranteed on the deploy machine, while
reportlab is a pure-Python wheel that installs cleanly and gives exact control
over headers/footers/page numbers. python-docx is pure-Python too (lxml wheel).
"""

from __future__ import annotations

import html
from io import BytesIO
from typing import Any

# Design tokens — mirrors the frontend palette (app/globals.css).
NAVY_950 = "#0A2E6E"
NAVY_800 = "#14408D"
NAVY_600 = "#3E6CB8"
BODY_INK = "#1F2937"
SURFACE = "#F7F9FC"

BRAND_LINE = "MERIDIAN  ·  AI Governance Assessment Brief"


def _esc(text: str) -> str:
    return html.escape(text or "", quote=False)


# ── DOCX ──────────────────────────────────────────────────────────────────


def _add_page_number_field(paragraph) -> None:
    """Insert a PAGE field into a footer paragraph (python-docx has no native
    API for page numbers, so the OOXML field is written directly)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def render_docx(brief: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    navy = RGBColor(0x0A, 0x2E, 0x6E)
    navy_800 = RGBColor(0x14, 0x40, 0x8D)
    ink = RGBColor(0x1F, 0x29, 0x37)

    doc = Document()

    # Base typography.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = ink
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # Header (brand).
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run(BRAND_LINE)
    run.font.size = Pt(8)
    run.font.color.rgb = navy
    run.font.bold = True

    # Footer (page number only — the generated-at stamp is deliberately
    # excluded from the exported document).
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_prefix = footer_p.add_run("Page ")
    page_prefix.font.size = Pt(8)
    page_prefix.font.color.rgb = navy_800
    _add_page_number_field(footer_p)
    for r in footer_p.runs:
        r.font.size = Pt(8)
        if r.font.color.rgb is None:
            r.font.color.rgb = navy_800

    s = brief["sections"]

    # Title block.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f"{brief.get('country', '')} — {brief.get('policy_title', '')}")
    tr.font.size = Pt(17)
    tr.font.bold = True
    tr.font.color.rgb = navy

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("AI Governance Assessment Brief")
    sr.font.size = Pt(12)
    sr.font.color.rgb = navy_800

    # (The generated-at / dimension-count stamp is on the on-screen card
    # only — deliberately NOT included in the exported document.)

    def heading(text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.color.rgb = navy if level == 1 else navy_800
            r.font.name = "Calibri"
        h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        h.paragraph_format.space_after = Pt(4)

    def body(text: str) -> None:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    def bullet(text: str) -> None:
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    # ── Sections ──────────────────────────────────────────────────────
    heading("EXECUTIVE SUMMARY")
    body(s["executive_summary"])

    heading("KEY FINDINGS")
    heading("Areas of Strength", level=2)
    if s["areas_of_strength"]:
        for b in s["areas_of_strength"]:
            bullet(b)
    else:
        body("None identified.")

    heading("Areas Requiring Attention", level=2)
    if s["areas_requiring_attention"]:
        for b in s["areas_requiring_attention"]:
            bullet(b)
    else:
        body("None identified.")

    heading("RISK OVERVIEW")
    body(s["risk_overview"]["paragraph"])

    # Deterministic depth sections (dimension detail, roadmap, evidence).
    # Rendered from stored analysis, so the exported document carries the same
    # substance as the on-screen brief rather than a shorter summary of it.
    rows = s.get("dimension_assessment") or []
    if rows:
        heading("DIMENSION ASSESSMENT")
        for r in rows:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            label = f"{r['dimension']} — {r['coverage']}"
            if r.get("maturity"):
                label += f" · {r['maturity']}"
            run = p.add_run(label)
            run.font.bold = True
            run.font.color.rgb = ink
            if r.get("basis"):
                body(r["basis"])
            if r.get("absent_mechanisms"):
                bullet("Not addressed: " + ", ".join(r["absent_mechanisms"]))

    heading("PRIORITY RECOMMENDATIONS")
    recs = s["priority_recommendations"]
    if recs:
        for i, r in enumerate(recs, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            nr = p.add_run(f"{i}. {r['recommendation']}")
            nr.font.bold = True
            nr.font.color.rgb = ink
            if r.get("rationale"):
                rr = p.add_run(f" — {r['rationale']}")
                rr.font.color.rgb = ink
    else:
        body("No critical gaps identified — no priority actions required.")

    roadmap = s.get("implementation_roadmap") or []
    if roadmap:
        heading("IMPLEMENTATION ROADMAP")
        for item in roadmap:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{item['dimension']} ({item['coverage']})")
            run.font.bold = True
            run.font.color.rgb = ink
            if item.get("responsible_agency"):
                body(f"Responsible body: {item['responsible_agency']}")
            for ph in item["phases"]:
                label = ph["phase"] or "Phase"
                if ph.get("timeline"):
                    label += f" · {ph['timeline']}"
                body(f"{label} — {ph.get('objective', '')}")
                for st in ph["steps"]:
                    bullet(st)
            for mc in item.get("monitoring") or []:
                bullet(f"Monitor: {mc}")

    ev = s.get("evidence_base") or {}
    if ev.get("citations_total"):
        heading("EVIDENCE BASE")
        body(
            f"{ev['citations_verified']} of {ev['citations_total']} citations were "
            "verified against their source passage."
        )
        for q in ev.get("representative_quotes") or []:
            bullet(f"{q['dimension']} — \u201c{q['quote']}\u201d")

    if s.get("relevant_precedent"):
        heading("RELEVANT PRECEDENT")
        body(s["relevant_precedent"])

    heading("SCOPE & METHODOLOGY")
    for para in s["scope_and_methodology"].split("\n\n"):
        body(para)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── PDF ──────────────────────────────────────────────────────────────────


def render_pdf(brief: dict[str, Any]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    s = brief["sections"]
    buffer = BytesIO()

    def _footer(canvas, doc_obj) -> None:
        # Brand left, page number right — the generated-at stamp is
        # deliberately excluded from the exported document.
        canvas.saveState()
        canvas.setFont("Helvetica-Bold", 7.5)
        canvas.setFillColor(colors.HexColor(NAVY_950))
        canvas.drawString(18 * mm, 12 * mm, BRAND_LINE)
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor(NAVY_800))
        canvas.drawRightString(A4[0] - 18 * mm, 12 * mm, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    title_style = ParagraphStyle(
        "Title",
        fontName="Helvetica-Bold",
        fontSize=16.5,
        leading=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor(NAVY_950),
        spaceAfter=2,
    )
    sub_style = ParagraphStyle(
        "Sub",
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor(NAVY_800),
        spaceAfter=2,
    )
    h1_style = ParagraphStyle(
        "H1",
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        spaceBefore=12,
        spaceAfter=4,
        textColor=colors.HexColor(NAVY_950),
    )
    h2_style = ParagraphStyle(
        "H2",
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        spaceBefore=8,
        spaceAfter=3,
        textColor=colors.HexColor(NAVY_800),
    )
    body_style = ParagraphStyle(
        "Body",
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        spaceAfter=6,
        textColor=colors.HexColor(BODY_INK),
        alignment=TA_LEFT,
    )

    story: list[Any] = []
    story.append(
        Paragraph(
            _esc(f"{brief.get('country', '')} — {brief.get('policy_title', '')}"), title_style
        )
    )
    story.append(Paragraph("AI Governance Assessment Brief", sub_style))
    story.append(Spacer(1, 4))

    def _h1(text: str) -> None:
        story.append(Paragraph(_esc(text), h1_style))

    def _h2(text: str) -> None:
        story.append(Paragraph(_esc(text), h2_style))

    def _body(text: str) -> None:
        story.append(Paragraph(_esc(text), body_style))

    def _bullets(items: list[str]) -> None:
        story.append(
            ListFlowable(
                [ListItem(Paragraph(_esc(item), body_style), leftIndent=14) for item in items],
                bulletType="bullet",
                start="•",
                leftIndent=12,
            )
        )
        story.append(Spacer(1, 2))

    _h1("EXECUTIVE SUMMARY")
    _body(s["executive_summary"])

    _h1("KEY FINDINGS")
    _h2("Areas of Strength")
    if s["areas_of_strength"]:
        _bullets(s["areas_of_strength"])
    else:
        _body("None identified.")
    _h2("Areas Requiring Attention")
    if s["areas_requiring_attention"]:
        _bullets(s["areas_requiring_attention"])
    else:
        _body("None identified.")

    _h1("RISK OVERVIEW")
    _body(s["risk_overview"]["paragraph"])

    rows = s.get("dimension_assessment") or []
    if rows:
        _h1("DIMENSION ASSESSMENT")
        for r in rows:
            label = f"{r['dimension']} — {r['coverage']}"
            if r.get("maturity"):
                label += f" · {r['maturity']}"
            _h2(label)
            if r.get("basis"):
                _body(r["basis"])
            if r.get("absent_mechanisms"):
                _bullets(["Not addressed: " + ", ".join(r["absent_mechanisms"])])

    _h1("PRIORITY RECOMMENDATIONS")
    recs = s["priority_recommendations"]
    if recs:
        items = [
            f"<b>{_esc(r['recommendation'])}</b>"
            + (f" — {_esc(r['rationale'])}" if r.get("rationale") else "")
            for r in recs
        ]
        story.append(
            ListFlowable(
                [ListItem(Paragraph(item, body_style), leftIndent=14) for item in items],
                bulletType="1",
                leftIndent=12,
            )
        )
        story.append(Spacer(1, 2))
    else:
        _body("No critical gaps identified — no priority actions required.")

    roadmap = s.get("implementation_roadmap") or []
    if roadmap:
        _h1("IMPLEMENTATION ROADMAP")
        for item in roadmap:
            _h2(f"{item['dimension']} ({item['coverage']})")
            if item.get("responsible_agency"):
                _body(f"Responsible body: {item['responsible_agency']}")
            for ph in item["phases"]:
                label = ph["phase"] or "Phase"
                if ph.get("timeline"):
                    label += f" \u00b7 {ph['timeline']}"
                _body(f"{label} — {ph.get('objective', '')}")
                _bullets(ph["steps"])
            if item.get("monitoring"):
                _bullets([f"Monitor: {mc}" for mc in item["monitoring"]])

    ev = s.get("evidence_base") or {}
    if ev.get("citations_total"):
        _h1("EVIDENCE BASE")
        _body(
            f"{ev['citations_verified']} of {ev['citations_total']} citations were "
            "verified against their source passage."
        )
        if ev.get("representative_quotes"):
            _bullets(
                [
                    f"{q['dimension']} — \u201c{q['quote']}\u201d"
                    for q in ev["representative_quotes"]
                ]
            )

    if s.get("relevant_precedent"):
        _h1("RELEVANT PRECEDENT")
        _body(s["relevant_precedent"])

    _h1("SCOPE & METHODOLOGY")
    for para in s["scope_and_methodology"].split("\n\n"):
        _body(para)

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=18 * mm,
        title=f"AI Governance Assessment Brief — {brief.get('country', '')}",
        author="Meridian",
    )
    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buffer.getvalue()
